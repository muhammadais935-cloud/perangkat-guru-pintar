import streamlit as st
import google.generativeai as genai
from docx import Document
from fpdf import FPDF
import io

st.set_page_config(page_title="Perangkat Guru Pintar AI", page_icon="🎓", layout="wide")

st.title("🎓 Perangkat Guru Pintar (AI-Powered)")
st.caption("Asisten Digital Berbasis Gemini AI dengan Fitur Ekspor Word (DOCX) & PDF")

# Sidebar - Pengaturan API Key & Navigasi
st.sidebar.header("⚙️ Pengaturan AI")
api_key = st.sidebar.text_input("Masukkan Gemini API Key:", type="password", help="Dapatkan API key gratis di Google AI Studio (aistudio.google.com)")

if api_key:
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel("gemini-1.5-flash")

menu = st.sidebar.radio("Pilih Fitur:", ["Generator Modul Ajar", "Pembuat Soal Kuis", "Rubrik Penilaian"])

# Helper Function: Generate DOCX
def create_docx(title, content):
    doc = Document()
    doc.add_heading(title, level=1)
    
    # Split by lines and add paragraphs
    for line in content.split('\n'):
        if line.startswith('# '):
            doc.add_heading(line.replace('# ', ''), level=1)
        elif line.startswith('## '):
            doc.add_heading(line.replace('## ', ''), level=2)
        elif line.startswith('### '):
            doc.add_heading(line.replace('### ', ''), level=3)
        elif line.startswith('* ') or line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        else:
            doc.add_paragraph(line)
            
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# Helper Function: Generate PDF
def create_pdf(title, content):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    
    # Title
    pdf.set_font("Helvetica", style="B", size=16)
    pdf.cell(0, 10, txt=title.encode('latin-1', 'replace').decode('latin-1'), ln=True, align="C")
    pdf.ln(5)
    
    # Content Body
    pdf.set_font("Helvetica", size=10)
    lines = content.split('\n')
    for line in lines:
        # Clean up markdown markers for simple PDF output
        clean_line = line.replace('**', '').replace('*', '').replace('#', '').strip()
        if clean_line:
            pdf.multi_cell(0, 6, txt=clean_line.encode('latin-1', 'replace').decode('latin-1'))
            pdf.ln(1)
            
    pdf_out = io.BytesIO()
    pdf_bytes = pdf.output()
    pdf_out.write(pdf_bytes)
    pdf_out.seek(0)
    return pdf_out

# ---------------- FITUR 1: MODUL AJAR ----------------
if menu == "Generator Modul Ajar":
    st.header("📋 Generator Modul Ajar AI")
    col1, col2 = st.columns(2)
    with col1:
        mapel = st.text_input("Mata Pelajaran:", "Informatika")
        kelas = st.selectbox("Fase / Kelas:", ["Fase A (Kelas 1-2)", "Fase B (Kelas 3-4)", "Fase C (Kelas 5-6)", "Fase D (Kelas 7-9)", "Fase E/F (Kelas 10-12)"])
    with col2:
        topik = st.text_input("Materi / Topik Utama:", "Berpikir Komputasional")
        alokasi = st.text_input("Alokasi Waktu:", "2 x 45 Menit")

    tujuan = st.text_area("Tujuan Pembelajaran (TP):", "Siswa mampu mengidentifikasi pola dalam penyelesaian masalah sehari-hari.")

    if st.button("Generate Modul Ajar"):
        if not api_key:
            st.error("Masukkan Google Gemini API Key terlebih dahulu di menu samping (sidebar).")
        else:
            with st.spinner("AI sedang menyusun Modul Ajar lengkap..."):
                prompt = f"""
                Buatkan draf Modul Ajar Kurikulum Merdeka yang sistematis untuk:
                - Mata Pelajaran: {mapel}
                - Fase/Kelas: {kelas}
                - Topik: {topik}
                - Alokasi Waktu: {alokasi}
                - Tujuan Pembelajaran: {tujuan}

                Format output memuat:
                1. Informasi Umum & Profil Pelajar Pancasila
                2. Pertanyaan Pemantik
                3. Kegiatan Pembelajaran (Pendahuluan, Inti, Penutup)
                4. Asesmen (Formatif & Sumatif)
                """
                response = model.generate_content(prompt)
                st.session_state['modul_result'] = response.text
                st.session_state['modul_title'] = f"Modul Ajar {mapel} - {topik}"

    if 'modul_result' in st.session_state:
        st.subheader("📄 Hasil Modul Ajar")
        st.markdown(st.session_state['modul_result'])
        
        st.divider()
        st.subheader("📥 Unduh Dokumen")
        d_col1, d_col2 = st.columns(2)
        
        docx_file = create_docx(st.session_state['modul_title'], st.session_state['modul_result'])
        with d_col1:
            st.download_button(
                label="📄 Unduh Format Word (.docx)",
                data=docx_file,
                file_name=f"{st.session_state['modul_title']}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            
        pdf_file = create_pdf(st.session_state['modul_title'], st.session_state['modul_result'])
        with d_col2:
            st.download_button(
                label="📕 Unduh Format PDF (.pdf)",
                data=pdf_file,
                file_name=f"{st.session_state['modul_title']}.pdf",
                mime="application/pdf"
            )

# ---------------- FITUR 2: SOAL KUIS ----------------
elif menu == "Pembuat Soal Kuis":
    st.header("❓ Pembuat Soal Kuis AI")
    col1, col2 = st.columns(2)
    with col1:
        topik_soal = st.text_input("Topik Kuis:", "Ekosistem Lingkungan")
        jumlah_soal = st.slider("Jumlah Soal:", 1, 10, 5)
    with col2:
        tingkat = st.selectbox("Tingkat Kesulitan:", ["Mudah", "Sedang", "HOTS (Tinggi)"])
        tipe_soal = st.selectbox("Tipe Soal:", ["Pilihan Ganda (dengan opsi A-D)", "Esai / Uraian"])

    if st.button("Generate Soal"):
        if not api_key:
            st.error("Masukkan Google Gemini API Key terlebih dahulu di menu samping (sidebar).")
        else:
            with st.spinner("AI sedang merancang soal..."):
                prompt = f"""
                Buatkan {jumlah_soal} soal {tipe_soal} tingkat kesulitan {tingkat} dengan topik '{topik_soal}'.
                Sertakan Kunci Jawaban dan Pembahasan Ringkas di bagian akhir.
                """
                response = model.generate_content(prompt)
                st.session_state['kuis_result'] = response.text
                st.session_state['kuis_title'] = f"Soal Kuis {topik_soal}"

    if 'kuis_result' in st.session_state:
        st.subheader("📄 Hasil Soal Kuis")
        st.markdown(st.session_state['kuis_result'])
        
        st.divider()
        st.subheader("📥 Unduh Dokumen")
        d_col1, d_col2 = st.columns(2)
        
        docx_file = create_docx(st.session_state['kuis_title'], st.session_state['kuis_result'])
        with d_col1:
            st.download_button(
                label="📄 Unduh Format Word (.docx)",
                data=docx_file,
                file_name=f"{st.session_state['kuis_title']}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            
        pdf_file = create_pdf(st.session_state['kuis_title'], st.session_state['kuis_result'])
        with d_col2:
            st.download_button(
                label="📕 Unduh Format PDF (.pdf)",
                data=pdf_file,
                file_name=f"{st.session_state['kuis_title']}.pdf",
                mime="application/pdf"
            )

# ---------------- FITUR 3: RUBRIK PENILAIAN ----------------
elif menu == "Rubrik Penilaian":
    st.header("📊 Generator Rubrik Penilaian AI")
    nama_tugas = st.text_input("Nama Tugas / Proyek:", "Presentasi Kelompok - Daur Ulang Sampah")
    kriteria_khusus = st.text_area("Kriteria Tambahan (Opsional):", "Kreativitas media, Kerjasama tim, Penguasaan materi")

    if st.button("Generate Rubrik"):
        if not api_key:
            st.error("Masukkan Google Gemini API Key terlebih dahulu di menu samping (sidebar).")
        else:
            with st.spinner("AI sedang menyusun tabel rubrik..."):
                prompt = f"""
                Buatkan rubrik penilaian analitis berbentuk tabel Markdown untuk tugas: '{nama_tugas}'.
                Kriteria penilaian mencakup: {kriteria_khusus}.
                Gunakan skala: Sangat Baik (4), Baik (3), Cukup (2), dan Perlu Bimbingan (1) lengkap dengan deskriptor indikatornya.
                """
                response = model.generate_content(prompt)
                st.session_state['rubrik_result'] = response.text
                st.session_state['rubrik_title'] = f"Rubrik Penilaian - {nama_tugas}"

    if 'rubrik_result' in st.session_state:
        st.subheader("📄 Hasil Rubrik Penilaian")
        st.markdown(st.session_state['rubrik_result'])
        
        st.divider()
        st.subheader("📥 Unduh Dokumen")
        d_col1, d_col2 = st.columns(2)
        
        docx_file = create_docx(st.session_state['rubrik_title'], st.session_state['rubrik_result'])
        with d_col1:
            st.download_button(
                label="📄 Unduh Format Word (.docx)",
                data=docx_file,
                file_name=f"{st.session_state['rubrik_title']}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            
        pdf_file = create_pdf(st.session_state['rubrik_title'], st.session_state['rubrik_result'])
        with d_col2:
            st.download_button(
                label="📕 Unduh Format PDF (.pdf)",
                data=pdf_file,
                file_name=f"{st.session_state['rubrik_title']}.pdf",
                mime="application/pdf"
            )
